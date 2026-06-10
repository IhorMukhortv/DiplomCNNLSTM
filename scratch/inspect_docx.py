import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

def inspect():
    doc = Document("docs/thesis/diploma_work.docx")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    
    count = 0
    for idx, p in enumerate(doc.paragraphs):
        p_el = p._p
        has_math = p_el.find("{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath") is not None or \
                   p_el.find("{http://schemas.openxmlformats.org/officeDocument/2006/math}oMathPara") is not None
                   
        if has_math and p.text.strip():  # Only display equations that are not blank
            print(f"Paragraph {idx}:")
            print(f"  Text: {repr(p.text)}")
            print(f"  Alignment: {p.alignment}")
            print(f"  First Line Indent: {p.paragraph_format.first_line_indent}")
            print(f"  Tabs count: {len(p.paragraph_format.tab_stops)}")
            for t_idx, tab in enumerate(p.paragraph_format.tab_stops):
                print(f"    Tab {t_idx}: position={tab.position.inches:.2f} inches, alignment={tab.alignment}")
            print("-" * 50)
            count += 1
            if count >= 20:
                break

if __name__ == "__main__":
    inspect()
