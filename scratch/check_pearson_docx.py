import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document("docs/thesis/diploma_work.docx")
print("Total paragraphs:", len(doc.paragraphs))
for idx in range(60, 90):
    if idx < len(doc.paragraphs):
        p = doc.paragraphs[idx]
        print(f"Paragraph {idx}: {repr(p.text)}")
