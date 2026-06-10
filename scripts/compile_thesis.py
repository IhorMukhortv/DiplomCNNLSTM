import os
import re
import sys
import shutil
import subprocess
import logging
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Налаштування логування
logging.basicConfig(level=logging.DEBUG, format="%(asctime)s - %(name)s - %(levelname)s - %(message)s")
logger = logging.getLogger("compile_thesis")

# Послідовність файлів розділів дипломної роботи
CHAPTER_FILES = [
    "docs/thesis/introduction.md",
    "docs/thesis/chapter1_data_collection_analysis.md",
    "docs/thesis/chapter2_theoretical_background.md",
    "docs/thesis/chapter3_model_design.md",
    "docs/thesis/chapter4_infrastructure.md",
    "docs/thesis/chapter5_training_evaluation.md",
    "docs/thesis/chapter6_online_learning.md",
    "docs/thesis/chapter7_practical_implementation.md",
    "docs/thesis/chapter8_economic_analysis.md",
    "docs/thesis/conclusions.md",
    "docs/thesis/references.md"
]

OUTPUT_DOCX = "docs/thesis/diploma_work.docx"
TEMP_MD = "docs/thesis/temp_thesis.md"

PAGE_BREAK = "\n\n```{=openxml}\n<w:p><w:r><w:br w:type=\"page\"/></w:r></w:p>\n```\n\n"

def find_pandoc():
    """Знаходить виконуваний файл Pandoc у системі."""
    if shutil.which("pandoc"):
        return "pandoc"
    if sys.platform == "win32":
        local_appdata = os.environ.get("LOCALAPPDATA", "")
        if local_appdata:
            fallback_path = os.path.join(local_appdata, "Pandoc", "pandoc.exe")
            if os.path.exists(fallback_path):
                return fallback_path
    return None

def clean_text_backslashes(text):
    """Очищає екрановані символи розмітки у тексті."""
    text = text.replace('\\%', '%')
    text = text.replace('\\_', '_')
    text = text.replace('\\#', '#')
    text = text.replace('\\$', '$')
    text = text.replace('\\&', '&')
    text = text.replace('\\{', '{')
    text = text.replace('\\}', '}')
    return text

def clean_latex_formula(text):
    """Спрощене очищення для тестових викликів (для сумісності з тестами)."""
    text = re.sub(r'\\qquad.*$', '', text).strip()
    text = re.sub(r'\\eqno.*$', '', text).strip()
    text = text.replace('\\frac', '/').replace('\\sum', '∑')
    return text

def parse_math_to_runs(text):
    """Спрощений розбір для сумісності з тестами."""
    runs = []
    i = 0
    n = len(text)
    while i < n:
        char = text[i]
        if char == '_':
            i += 1
            start = i
            while i < n and text[i].isalnum():
                i += 1
            runs.append((text[start:i], True, False))
        else:
            start = i
            while i < n and text[i] != '_':
                i += 1
            runs.append((text[start:i], False, False))
    return runs

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Налаштування внутрішніх відступів у комірках таблиці."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table):
    """Налаштування академічних меж таблиці (ДСТУ/академічний стиль)."""
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            '  <w:top w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
            '  <w:left w:val="none"/>'
            '  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
            '  <w:right w:val="none"/>'
            '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="888888"/>'
            '  <w:insideV w:val="none"/>'
            '</w:tblBorders>'
        )
        tblPr[0].append(borders)

def format_heading_style(style, font_size=14, bold=True, italic=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=12):
    """Форматування стилю заголовків."""
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(font_size)
    font.bold = bold
    font.italic = italic
    font.color.rgb = RGBColor(0, 0, 0)
    
    if hasattr(style, 'paragraph_format'):
        p_format = style.paragraph_format
        p_format.line_spacing = 1.5
        p_format.space_before = Pt(space_before)
        p_format.space_after = Pt(space_after)
        p_format.first_line_indent = Inches(0)
        p_format.alignment = alignment

def paragraph_has_image(p):
    """Перевіряє, чи містить параграф зображення."""
    p_el = p._p
    if p_el.find(qn('w:drawing')) is not None:
        return True
    if p_el.find(qn('pic:pic')) is not None:
        return True
    return False

def find_math_element(p):
    """Шукає математичні блоки в параграфі."""
    p_el = p._p
    math_para = p_el.find(qn('m:oMathPara'))
    if math_para is not None:
        return math_para
    math = p_el.find(qn('m:oMath'))
    if math is not None:
        return math
    return None

def preprocess_markdown(content):
    """Об'єднує контент і готує блок-формули та заголовки до Pandoc-компіляції."""
    # Видаляємо горизонтальні роздільники ---
    content = re.sub(r'^\s*-{3,}\s*$', '', content, flags=re.MULTILINE)

    # Очищаємо провідні косі риски у шляхах до зображень для локального рендерингу Pandoc
    content = re.sub(r'!\[(.*?)\]\(/+(.*?)\)', r'![\1](\2)', content)
    
    # Заголовки розділів (#) робимо великими літерами за вимогами ДСТУ
    def uppercase_headers(match):
        return f"# {match.group(1).upper()}"
    content = re.sub(r'^#\s+(.+)$', uppercase_headers, content, flags=re.MULTILINE)
    
    # Вилучаємо номери формул і переформатовуємо їх під пост-процесор
    def replace_equation(match):
        formula_content = match.group(1).strip()
        block_text = match.group(0)
        
        # Шукаємо номер формули
        num_match = re.search(r'\(\s*(\d+\.\d+|\d+)\s*\)', block_text)
        eq_num = num_match.group(0) if num_match else ""
        
        formula_clean = re.sub(r'\\qquad.*$', '', formula_content).strip()
        formula_clean = re.sub(r'\\eqno.*$', '', formula_clean).strip()
        if eq_num:
            formula_clean = formula_clean.replace(eq_num, '').strip()
            
        if eq_num:
            return f"\n\n$${formula_clean}$$\n\n[EQNO: {eq_num}]\n\n"
        else:
            return f"\n\n$${formula_clean}$$\n\n"
            
    content = re.sub(r'\$\$(.*?)\$\$', replace_equation, content, flags=re.DOTALL)
    return content

def post_process_docx(docx_path):
    """Пост-обробка згенерованого Word файлу за допомогою python-docx."""
    logger.info("Запуск пост-процесора для налаштування стилів за ДСТУ...")
    doc = Document(docx_path)
    
    # Налаштування полів сторінки
    for section in doc.sections:
        section.top_margin = Inches(0.787)      # 20 мм
        section.bottom_margin = Inches(0.787)   # 20 мм
        section.left_margin = Inches(0.984)     # 25 мм
        section.right_margin = Inches(0.59)      # 15 мм

    # Налаштування стилю Normal (основний текст)
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.color.rgb = RGBColor(0, 0, 0)
    p_format = style_normal.paragraph_format
    p_format.line_spacing = 1.5
    p_format.space_after = Pt(6)
    p_format.space_before = Pt(0)
    p_format.first_line_indent = Inches(0.5)
    p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Налаштування заголовків (Heading 1 - Heading 6)
    for style in doc.styles:
        name = style.name
        if name.startswith('Heading '):
            try:
                level = int(name.split(' ')[1])
                if level == 1:
                    format_heading_style(style, font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=18)
                elif level == 2:
                    format_heading_style(style, font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=12)
                elif level == 3:
                    format_heading_style(style, font_size=14, bold=True, italic=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6)
                else:
                    # Для Heading 4, Heading 5, Heading 6
                    format_heading_style(style, font_size=14, bold=True, italic=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=6)
            except (ValueError, IndexError):
                pass

    # Увімкнення автоматичного переносу слів (Auto-hyphenation)
    settings = doc.settings
    element = settings.element
    auto_hyphen = element.find(qn('w:autoHyphenation'))
    if auto_hyphen is None:
        auto_hyphen = OxmlElement('w:autoHyphenation')
        auto_hyphen.set(qn('w:val'), 'true')
        after_elements = [
            'w:consecutiveHyphenLimit', 'w:hyphenationZone', 'w:doNotHyphenateCaps',
            'w:showXMLTags', 'w:compatibility', 'w:rsids', 'w:mathPr', 'w:compatSetting'
        ]
        inserted = False
        for child_tag in after_elements:
            child = element.find(qn(child_tag))
            if child is not None:
                idx = element.index(child)
                element.insert(idx, auto_hyphen)
                inserted = True
                break
        if not inserted:
            element.append(auto_hyphen)

    # Налаштування стилів списків
    for style_name in ['List Bullet', 'List Number']:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = 'Times New Roman'
            style.font.size = Pt(14)
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.paragraph_format.line_spacing = 1.5
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(3)

    # Пошук та форматування блок-формул із номерами праворуч
    paragraphs = list(doc.paragraphs)
    i = 0
    while i < len(paragraphs):
        p = paragraphs[i]
        if "[EQNO:" in p.text:
            num_match = re.search(r'\[EQNO:\s*(\(\d+\.\d+\))\s*\]', p.text)
            if num_match:
                eq_num = num_match.group(1)
                
                # Шукаємо найближчий попередній параграф з математикою
                eq_p = None
                for offset in [1, 2]:
                    if i - offset >= 0:
                        cand = paragraphs[i - offset]
                        if find_math_element(cand) is not None:
                            eq_p = cand
                            break
                            
                if eq_p is not None:
                    p_el = eq_p._p
                    # Вилучаємо вкладений oMathPara і переміщуємо oMath безпосередньо під w:p
                    math_para = p_el.find(qn('m:oMathPara'))
                    if math_para is not None:
                        math = math_para.find(qn('m:oMath'))
                        if math is not None:
                            idx = p_el.index(math_para)
                            p_el.remove(math_para)
                            p_el.insert(idx, math)
                    
                    # Додаємо табуляцію на початку
                    r_el = OxmlElement('w:r')
                    tab_el = OxmlElement('w:tab')
                    r_el.append(tab_el)
                    pPr = p_el.find(qn('w:pPr'))
                    if pPr is not None:
                        idx = p_el.index(pPr)
                        p_el.insert(idx + 1, r_el)
                    else:
                        p_el.insert(0, r_el)
                        
                    # Додаємо табуляцію та номер формули наприкінці
                    r = eq_p.add_run('\t' + eq_num)
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(14)
                    r.font.color.rgb = RGBColor(0, 0, 0)
                    
                    # Налаштовуємо Tab Stops
                    eq_p.paragraph_format.tab_stops.add_tab_stop(Inches(3.35), alignment=WD_TAB_ALIGNMENT.CENTER)
                    eq_p.paragraph_format.tab_stops.add_tab_stop(Inches(6.70), alignment=WD_TAB_ALIGNMENT.RIGHT)
                    eq_p.paragraph_format.first_line_indent = Inches(0)
                    eq_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    eq_p.paragraph_format.space_before = Pt(6)
                    eq_p.paragraph_format.space_after = Pt(6)
                    eq_p.paragraph_format.line_spacing = 1.5
                    
            # Видаляємо маркерний параграф
            p_el = p._p
            p_el.getparent().remove(p_el)
        i += 1

    # Центрування зображень та форматування підписів до них
    for p in doc.paragraphs:
        if paragraph_has_image(p):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Inches(0)
            
        text_clean = p.text.strip()
        if text_clean.startswith("де "):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Inches(0.5)

        if text_clean.startswith("Рисунок") or text_clean.startswith("Таблиця") or text_clean.startswith("Рисунок."):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(12)
            for run in p.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                run.italic = True

    # Форматування нативних таблиць
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table)
        for row in table.rows:
            for cell in row.cells:
                set_cell_margins(cell)
                for cell_p in cell.paragraphs:
                    cell_p.paragraph_format.line_spacing = 1.15
                    cell_p.paragraph_format.space_before = Pt(2)
                    cell_p.paragraph_format.space_after = Pt(2)
                    cell_p.paragraph_format.first_line_indent = Inches(0)
                    for run in cell_p.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)
                        
    doc.save(docx_path)
    logger.info("Пост-обробка успішно завершена.")

def compile_markdown_to_docx():
    logger.info("Початок компіляції дипломної роботи через Pandoc...")
    pandoc_bin = find_pandoc()
    if not pandoc_bin:
        logger.error("Pandoc не знайдено! Встановіть його за допомогою: winget install JohnMacFarlane.Pandoc")
        sys.exit(1)
        
    logger.debug(f"Використовується виконуваний файл Pandoc: {pandoc_bin}")
    
    # Об'єднуємо глави у великий Markdown
    combined_content = ""
    first_chapter = True
    
    for file_path in CHAPTER_FILES:
        if not os.path.exists(file_path):
            logger.warning(f"Файл {file_path} не знайдено, пропуск.")
            continue
            
        logger.debug(f"Зчитування {file_path}...")
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
            
        if not first_chapter:
            combined_content += PAGE_BREAK
        else:
            first_chapter = False
            
        combined_content += content + "\n\n"
        
    # Препроцесинг контенту перед згортанням в Pandoc
    combined_content = preprocess_markdown(combined_content)
    
    # Записуємо у тимчасовий файл
    with open(TEMP_MD, "w", encoding="utf-8") as f:
        f.write(combined_content)
        
    try:
        # Компіляція через Pandoc
        logger.info("Запуск конвертації Markdown в DOCX через Pandoc...")
        cmd = [pandoc_bin, TEMP_MD, "-o", OUTPUT_DOCX]
        subprocess.run(cmd, check=True)
        
        # Пост-обробка згенерованого файлу
        post_process_docx(OUTPUT_DOCX)
        logger.info(f"Скомпільовано успішно: {OUTPUT_DOCX}")
    finally:
        # Видаляємо тимчасовий файл
        if os.path.exists(TEMP_MD):
            os.remove(TEMP_MD)

if __name__ == "__main__":
    compile_markdown_to_docx()
